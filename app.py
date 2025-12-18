import os
import logging
from contextlib import asynccontextmanager
from typing import List, Optional

import fitz  # PyMuPDF
import docx
import google.generativeai as genai
import uvicorn
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# Загрузка переменных окружения
load_dotenv()

# Конфигурация
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    logger.error("GOOGLE_API_KEY не найден в переменных окружения")
    raise ValueError("GOOGLE_API_KEY обязателен")

# Используем актуальную модель (1.5 Flash - быстрая и дешевая/бесплатная)
MODEL_NAME = os.getenv("GEMINI_MODEL", "gemini-flash-lite-latest")
DOCX_PATH = "data2.docx"
PDF_PATH = "data.pdf"

genai.configure(api_key=GOOGLE_API_KEY)

# --- Функции извлечения текста ---

def extract_text_from_pdf(pdf_path: str) -> str:
    text = ""
    try:
        with fitz.open(pdf_path) as doc:
            for page in doc:
                text += page.get_text()
        return text
    except Exception as e:
        logger.error(f"Ошибка при чтении PDF {pdf_path}: {e}")
        return ""

def extract_text_from_docx(docx_path: str) -> str:
    text = []
    try:
        doc = docx.Document(docx_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    text.append(" | ".join(row_text))
        
        return "\n".join(text)
    except Exception as e:
        logger.error(f"Ошибка при чтении DOCX {docx_path}: {e}")
        return ""

def load_resume_content() -> str:
    """Пытается загрузить резюме из DOCX, затем из PDF."""
    if os.path.exists(DOCX_PATH):
        content = extract_text_from_docx(DOCX_PATH)
        if content:
            logger.info(f"Загружено из DOCX ({len(content)} символов)")
            return content
            
    if os.path.exists(PDF_PATH):
        content = extract_text_from_pdf(PDF_PATH)
        if content:
            logger.info(f"Загружено из PDF ({len(content)} символов)")
            return content
            
    logger.warning("Файлы резюме не найдены.")
    return ""

# --- Жизненный цикл приложения ---

@asynccontextmanager
async def lifespan(app: FastAPI):
    # 1. Загружаем текст резюме
    resume_text = load_resume_content()
    
    if not resume_text:
        logger.error("Критическая ошибка: Текст резюме пуст.")
        # Можно либо остановить приложение, либо разрешить работу с заглушкой
    
    # 2. Формируем системный промпт (System Instruction)
    # Это задает поведение модели один раз при инициализации
    system_instruction = f"""
    Ты AI-ассистент рекрутера/посетителя сайта. Твоя цель — отвечать на вопросы о резюме кандидата (Колесников Максим).
    
    ПРАВИЛА:
    1. Отвечай ТОЛЬКО на основе предоставленного ниже текста резюме. Не выдумывай факты.
    2. Отвечай в третьем лице (например: "Максим работал...", "Он знает..."). Не говори "Я".
    3. Будь вежлив, краток и профессионален.
    4. Если ответа нет в резюме, скажи: "К сожалению, в резюме нет этой информации".
    5. Текст резюме:
    
    {resume_text}
    """

    # 3. Инициализируем модель и сохраняем в state
    try:
        # В Gemini 1.5 system_instruction передается здесь
        app.state.model = genai.GenerativeModel(
            model_name=MODEL_NAME,
            system_instruction=system_instruction
        )
        # Создаем чат-сессию (если хотим помнить контекст беседы, 
        # но для REST API часто проще генерировать ответ заново, 
        # либо хранить историю на клиенте. Здесь используем простой generate_content)
        logger.info(f"Модель {MODEL_NAME} инициализирована успешно.")
    except Exception as e:
        logger.error(f"Ошибка инициализации Gemini: {e}")
    
    yield
    
    # Очистка ресурсов (если нужно)
    logger.info("Приложение останавливается.")

# --- Инициализация FastAPI ---

app = FastAPI(title="Resume Chatbot API", lifespan=lifespan)

# Настройка CORS
origins = [
    "https://maks-mk.github.io",
    "http://localhost:8000",
    "http://127.0.0.1:8000",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins, # Лучше избегать "*", если возможно
    allow_credentials=True,
    allow_methods=["POST", "GET", "OPTIONS"],
    allow_headers=["*"],
)

class ChatRequest(BaseModel):
    message: str

# --- Эндпоинты ---

@app.post("/chat")
async def chat(request: ChatRequest, req: Request):
    if not hasattr(req.app.state, "model"):
        raise HTTPException(status_code=503, detail="AI сервис временно недоступен")

    try:
        # Отправляем сообщение пользователя.
        # Системная инструкция уже "вшита" в модель при инициализации.
        response = req.app.state.model.generate_content(request.message)
        
        return {"response": response.text}
    
    except Exception as e:
        logger.error(f"Ошибка при генерации ответа: {e}")
        raise HTTPException(status_code=500, detail="Ошибка обработки запроса к AI")

@app.api_route("/health", methods=["GET", "HEAD"])
async def health_check():
    return {"status": "ok", "model": MODEL_NAME}
    
# Статика (только для локальной разработки)
if os.getenv("ENVIRONMENT", "development") != "production":
    if os.path.exists("index.html"): # Проверка наличия фронтенда
        app.mount("/", StaticFiles(directory=".", html=True), name="static")

if __name__ == "__main__":
    port = int(os.getenv("PORT", 8000))
    uvicorn.run("app:app", host="0.0.0.0", port=port, reload=True)